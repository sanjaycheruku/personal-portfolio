import os, re, json, tempfile, io
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text as pdf_extract_text
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ReportLab for PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER

app = Flask(__name__, static_folder="static", static_url_path="")
CORS(app)
UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}

# ─── SKILL TAXONOMY ───────────────────────────────────
SKILL_TAXONOMY = {
    "Programming Languages": [
        "python","java","javascript","typescript","c++","c#","go","rust","ruby",
        "php","swift","kotlin","scala","r","matlab","perl","bash"
    ],
    "Web & Frontend": [
        "html","css","react","angular","vue","next.js","svelte","tailwind",
        "bootstrap","jquery","graphql","rest api","restful","webpack","vite"
    ],
    "Backend & Databases": [
        "node.js","django","flask","spring","fastapi","express","postgresql",
        "mysql","mongodb","redis","sqlite","firebase","supabase","cassandra"
    ],
    "Cloud & DevOps": [
        "aws","azure","gcp","docker","kubernetes","ci/cd","terraform","ansible",
        "jenkins","github actions","linux","nginx","heroku","vercel"
    ],
    "Data & AI/ML": [
        "machine learning","deep learning","nlp","computer vision","tensorflow",
        "pytorch","scikit-learn","pandas","numpy","matplotlib","seaborn","tableau",
        "power bi","sql","spark","hadoop","data analysis","gen ai","generative ai",
        "llm","langchain","openai","hugging face"
    ],
    "Tools & Practices": [
        "git","github","jira","agile","scrum","kanban","devops","servicenow",
        "salesforce","sap","excel","powerpoint","figma","notion"
    ],
    "Soft Skills": [
        "leadership","communication","teamwork","problem solving","critical thinking",
        "project management","time management","mentoring","collaboration"
    ],
}
ALL_SKILLS = {skill: cat for cat, skills in SKILL_TAXONOMY.items() for skill in skills}
ATS_SECTIONS = ["experience","education","skills","projects","certifications",
                "summary","objective","achievements","internship"]
ACTION_VERBS = ["developed","designed","built","implemented","managed","led","created",
                "optimized","analyzed","collaborated","delivered","launched","automated",
                "reduced","increased","improved","architected","deployed","maintained"]

# ─── HELPERS ─────────────────────────────────────────
def allowed_file(fn): return "." in fn and fn.rsplit(".",1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(filepath, ext):
    if ext == "pdf":
        try: return pdf_extract_text(filepath)
        except Exception as e: return f"[PDF error: {e}]"
    elif ext == "docx":
        try:
            doc = docx.Document(filepath)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e: return f"[DOCX error: {e}]"
    else:
        with open(filepath,"r",encoding="utf-8",errors="ignore") as f: return f.read()

def extract_skills(text):
    text_lower = text.lower()
    found = {}
    for skill, cat in ALL_SKILLS.items():
        if re.search(r"\b" + re.escape(skill) + r"\b", text_lower):
            found.setdefault(cat, []).append(skill.title())
    return found

def compute_ats_score(text, job_desc=""):
    score, breakdown = 0, {}
    t = text.lower()

    # 1. Section structure (20pts)
    secs = [s for s in ATS_SECTIONS if s in t]
    ss = min(20, len(secs) * 3)
    breakdown["Section Structure"] = {"score": ss, "max": 20, "detail": f"Detected: {', '.join(secs[:5]) or 'none'}"}
    score += ss

    # 2. Keywords / Skills (25pts)
    found_skills = [s for s in ALL_SKILLS if re.search(r"\b"+re.escape(s)+r"\b", t)]
    ks = min(25, len(found_skills) * 2)
    breakdown["Skill Keywords"] = {"score": ks, "max": 25, "detail": f"{len(found_skills)} skills detected"}
    score += ks

    # 3. Contact info (10pts)
    cs = 0
    if re.search(r"[\w.+\-]+@[\w\-]+\.[a-z]{2,}", t): cs += 4
    if re.search(r"\+?\d[\d\s\-()\\.]{7,}", text): cs += 3
    if re.search(r"linkedin\.com", t): cs += 3
    breakdown["Contact Information"] = {"score": cs, "max": 10, "detail": "Email · Phone · LinkedIn"}
    score += cs

    # 4. Quantified achievements (20pts)
    nums = re.findall(r"\b\d+\s*(%|x|k|m|\+)?\b", t)
    verbs = [v for v in ACTION_VERBS if v in t]
    qs = min(20, len(nums)*2 + len(verbs))
    breakdown["Impact & Achievements"] = {"score": qs, "max": 20,
        "detail": f"{len(nums)} numbers, {len(verbs)} action verbs found"}
    score += qs

    # 5. Formatting & Length (10pts)
    wc = len(text.split())
    fs = 10 if 300 <= wc <= 800 else (7 if 200 <= wc < 300 or 800 < wc <= 1200 else 4)
    breakdown["Format & Length"] = {"score": fs, "max": 10, "detail": f"~{wc} words"}
    score += fs

    # 6. JD Match (15pts)
    if job_desc.strip():
        try:
            vec = TfidfVectorizer(stop_words="english")
            mat = vec.fit_transform([text, job_desc])
            sim = cosine_similarity(mat[0:1], mat[1:2])[0][0]
            jms = round(sim * 15)
        except: jms = 0
    else:
        jms = min(15, len(verbs) * 2)
    breakdown["JD Alignment"] = {"score": jms, "max": 15, "detail": "TF-IDF similarity" if job_desc.strip() else "Action verb heuristic"}
    score += jms

    return min(100, score), breakdown

def compute_job_match(resume_text, job_desc):
    if not job_desc.strip(): return 0.0, [], []
    try:
        jd_skills  = [s for s in ALL_SKILLS if re.search(r"\b"+re.escape(s)+r"\b", job_desc.lower())]
        res_skills = [s for s in ALL_SKILLS if re.search(r"\b"+re.escape(s)+r"\b", resume_text.lower())]
        matched  = [s for s in jd_skills if s in res_skills]
        missing  = [s for s in jd_skills if s not in res_skills]
        vec = TfidfVectorizer(stop_words="english")
        mat = vec.fit_transform([resume_text, job_desc])
        sim = round(cosine_similarity(mat[0:1], mat[1:2])[0][0] * 100, 1)
        return sim, matched, missing
    except: return 0.0, [], []

def generate_suggestions(score, breakdown, missing, skills_found):
    tips = []
    if breakdown.get("Section Structure",{}).get("score",0) < 14:
        tips.append("📋 Add clearly labelled sections: Experience, Education, Skills, Projects, Certifications.")
    if breakdown.get("Skill Keywords",{}).get("score",0) < 18:
        tips.append("🔑 Include more technical keywords (languages, frameworks, tools) matching your target role.")
    if breakdown.get("Contact Information",{}).get("score",0) < 8:
        tips.append("📞 Ensure your resume has a professional email, phone number, and LinkedIn URL.")
    if breakdown.get("Impact & Achievements",{}).get("score",0) < 14:
        tips.append("📊 Quantify achievements — e.g., 'Improved API latency by 40%', 'Led team of 8 engineers'.")
    if breakdown.get("Format & Length",{}).get("score",0) < 7:
        tips.append("📄 Aim for 300–800 words (1 page recommended for < 5 years of experience).")
    if missing:
        tips.append(f"⚠️  Add missing JD skills: {', '.join(s.title() for s in missing[:6])}.")
    if not skills_found:
        tips.append("🛠️  No skills section detected — add a dedicated Skills section with relevant technologies.")
    if score >= 80:
        tips.append("✅ Outstanding ATS score! You're in the top tier of candidates.")
    elif score >= 60:
        tips.append("⚡ Good score — a few targeted edits will elevate you significantly.")
    else:
        tips.append("🚀 Your resume needs keyword improvements and stronger impact statements.")
    return tips

# ─── PDF GENERATION ──────────────────────────────────
TEMPLATE_STYLES = {
    "modern":    {"primary": colors.HexColor("#6c63ff"), "accent": colors.HexColor("#a78bfa"), "font": "Helvetica"},
    "executive": {"primary": colors.HexColor("#1e3a5f"), "accent": colors.HexColor("#4a90d9"), "font": "Times-Roman"},
    "creative":  {"primary": colors.HexColor("#e85d04"), "accent": colors.HexColor("#f48c06"), "font": "Helvetica"},
    "minimal":   {"primary": colors.HexColor("#222222"), "accent": colors.HexColor("#555555"), "font": "Helvetica"},
}

def generate_pdf(data: dict, template: str = "modern") -> bytes:
    buf = io.BytesIO()
    ts  = TEMPLATE_STYLES.get(template, TEMPLATE_STYLES["modern"])
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=1.8*cm, bottomMargin=1.5*cm)
    story = []
    base  = ts["font"]
    bold  = base + "-Bold" if "Helvetica" in base else base + "-Bold"
    try: bold = base + "-Bold"
    except: bold = "Helvetica-Bold"

    def para(txt, size=10, bold_=False, color=colors.black, align=TA_LEFT, space_before=0, space_after=4):
        style = ParagraphStyle("x", fontName=bold if bold_ else base, fontSize=size,
                               textColor=color, alignment=align,
                               spaceBefore=space_before, spaceAfter=space_after,
                               leading=size*1.35)
        return Paragraph(txt or "", style)

    def hr(c): return HRFlowable(width="100%", thickness=1.5, color=c, spaceAfter=6)
    def section_title(title):
        story.append(Spacer(1, 0.3*cm))
        story.append(para(title.upper(), size=10, bold_=True, color=ts["primary"], space_before=8))
        story.append(hr(ts["accent"]))

    # Header
    name = data.get("name","Your Name")
    story.append(para(name, size=22, bold_=True, color=ts["primary"], align=TA_CENTER, space_after=2))
    contact_parts = [x for x in [data.get("email",""), data.get("phone",""), data.get("linkedin","")] if x]
    if contact_parts:
        story.append(para(" · ".join(contact_parts), size=9, color=colors.gray, align=TA_CENTER, space_after=8))
    story.append(hr(ts["primary"]))

    # Summary
    if data.get("summary","").strip():
        section_title("Professional Summary")
        story.append(para(data["summary"], size=10))

    # Experience
    exp_list = data.get("experience", [])
    if any(e.get("company") or e.get("role") for e in exp_list):
        section_title("Work Experience")
        for exp in exp_list:
            if not (exp.get("company") or exp.get("role")): continue
            row = f"<b>{exp.get('role','')}</b> — {exp.get('company','')}"
            period = exp.get("period","")
            tbl_data = [[Paragraph(row, ParagraphStyle("r", fontName=base, fontSize=10, leading=13)),
                         Paragraph(period, ParagraphStyle("p", fontName=base, fontSize=9,
                                                           textColor=colors.gray, alignment=1, leading=13))]]
            t = Table(tbl_data, colWidths=["70%","30%"])
            t.setStyle(TableStyle([("VALIGN","all",(0,0),(-1,-1),"TOP"),
                                   ("LEFTPADDING","all",(0,0),(-1,-1),0),
                                   ("RIGHTPADDING","all",(0,0),(-1,-1),0)]))
            story.append(t)
            if exp.get("description","").strip():
                for line in exp["description"].split("\n"):
                    line = line.strip()
                    if line:
                        story.append(para(f"• {line}", size=9, color=colors.HexColor("#333333"), space_after=2))
            story.append(Spacer(1, 0.2*cm))

    # Education
    edu_list = data.get("education", [])
    if any(e.get("institution") or e.get("degree") for e in edu_list):
        section_title("Education")
        for edu in edu_list:
            if not (edu.get("institution") or edu.get("degree")): continue
            row = f"<b>{edu.get('degree','')}</b> — {edu.get('institution','')}"
            period = edu.get("year","")
            tbl_data = [[Paragraph(row, ParagraphStyle("r2", fontName=base, fontSize=10, leading=13)),
                         Paragraph(period, ParagraphStyle("p2", fontName=base, fontSize=9,
                                                           textColor=colors.gray, alignment=1, leading=13))]]
            t = Table(tbl_data, colWidths=["70%","30%"])
            t.setStyle(TableStyle([("VALIGN","all",(0,0),(-1,-1),"TOP"),
                                   ("LEFTPADDING","all",(0,0),(-1,-1),0),
                                   ("RIGHTPADDING","all",(0,0),(-1,-1),0)]))
            story.append(t)
            if edu.get("grade","").strip():
                story.append(para(f"Grade/CGPA: {edu['grade']}", size=9, color=colors.gray))
            story.append(Spacer(1, 0.15*cm))

    # Skills
    skills_data = data.get("skills", {})
    if skills_data:
        section_title("Skills")
        for cat, skill_list in skills_data.items():
            if skill_list:
                story.append(para(f"<b>{cat}:</b> {', '.join(skill_list)}", size=9, space_after=3))

    # Projects
    proj_list = data.get("projects", [])
    if any(p.get("name") for p in proj_list):
        section_title("Projects")
        for proj in proj_list:
            if not proj.get("name"): continue
            story.append(para(f"<b>{proj.get('name','')}</b>  <font size='8' color='gray'>{proj.get('tech','')}</font>", size=10))
            if proj.get("description","").strip():
                story.append(para(proj["description"], size=9, color=colors.HexColor("#333333"), space_after=4))

    # Certifications
    cert_list = data.get("certifications", [])
    if cert_list:
        section_title("Certifications")
        for c in cert_list:
            if c.strip(): story.append(para(f"• {c}", size=9, space_after=2))

    doc.build(story)
    buf.seek(0)
    return buf.read()

def generate_docx(data: dict) -> bytes:
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = DocxDoc()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)

    def add_heading(text, level=1, color=(108,99,255)):
        h = doc.add_heading(text, level=level)
        r = h.runs[0]; r.font.color.rgb = RGBColor(*color); return h

    def add_para(text, bold=False, italic=False, color=None, size=10):
        p = doc.add_paragraph(); run = p.add_run(text)
        run.bold     = bold
        run.italic   = italic
        run.font.size = Pt(size)
        if color: run.font.color.rgb = RGBColor(*color)
        return p

    # Name & contact
    name_para = doc.add_paragraph()
    name_run   = name_para.add_run(data.get("name","Your Name"))
    name_run.bold      = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(108,99,255)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contacts = [x for x in [data.get("email",""), data.get("phone",""), data.get("linkedin","")] if x]
    if contacts:
        cp = doc.add_paragraph(" · ".join(contacts)); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs: r.font.size = Pt(9)

    doc.add_paragraph()

    if data.get("summary","").strip():
        add_heading("PROFESSIONAL SUMMARY", 2)
        add_para(data["summary"])

    exp_list = data.get("experience", [])
    if any(e.get("company") or e.get("role") for e in exp_list):
        add_heading("WORK EXPERIENCE", 2)
        for exp in exp_list:
            if not (exp.get("company") or exp.get("role")): continue
            add_para(f"{exp.get('role','')}  —  {exp.get('company','')}  ({exp.get('period','')})", bold=True)
            if exp.get("description","").strip():
                for line in exp["description"].split("\n"):
                    if line.strip(): add_para(f"• {line.strip()}")
            doc.add_paragraph()

    edu_list = data.get("education", [])
    if any(e.get("institution") or e.get("degree") for e in edu_list):
        add_heading("EDUCATION", 2)
        for edu in edu_list:
            if not (edu.get("institution") or edu.get("degree")): continue
            add_para(f"{edu.get('degree','')}  —  {edu.get('institution','')}  ({edu.get('year','')})", bold=True)
            if edu.get("grade","").strip(): add_para(f"Grade/CGPA: {edu['grade']}", italic=True)
            doc.add_paragraph()

    skills_data = data.get("skills", {})
    if skills_data:
        add_heading("SKILLS", 2)
        for cat, skill_list in skills_data.items():
            if skill_list: add_para(f"{cat}: {', '.join(skill_list)}", bold=False)

    proj_list = data.get("projects", [])
    if any(p.get("name") for p in proj_list):
        add_heading("PROJECTS", 2)
        for proj in proj_list:
            if not proj.get("name"): continue
            add_para(f"{proj.get('name','')}  |  {proj.get('tech','')}", bold=True)
            if proj.get("description","").strip(): add_para(proj["description"])
            doc.add_paragraph()

    cert_list = data.get("certifications", [])
    if cert_list:
        add_heading("CERTIFICATIONS", 2)
        for c in cert_list:
            if c.strip(): add_para(f"• {c}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ─── ROUTES ──────────────────────────────────────────
@app.route("/")
def index():
    return send_from_directory("static", "index.html")

@app.route("/analyze", methods=["POST"])
def analyze():
    if "resume" not in request.files:
        return jsonify({"error": "No resume file provided"}), 400
    file     = request.files["resume"]
    job_desc = request.form.get("job_description","")
    if not file.filename: return jsonify({"error":"No file selected"}), 400
    if not allowed_file(file.filename): return jsonify({"error":"Unsupported format. Use PDF, DOCX, or TXT."}), 400

    filename = secure_filename(file.filename)
    ext      = filename.rsplit(".",1)[1].lower()
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    try:
        resume_text = extract_text_from_file(filepath, ext)
        if len(resume_text.strip()) < 50:
            return jsonify({"error":"Could not extract enough text from the file."}), 400

        skills          = extract_skills(resume_text)
        score, brkdown  = compute_ats_score(resume_text, job_desc)
        jd_sim, matched, missing = compute_job_match(resume_text, job_desc)
        suggestions     = generate_suggestions(score, brkdown, missing, skills)
        wc              = len(resume_text.split())

        return jsonify({
            "ats_score":       score,
            "word_count":      wc,
            "skills":          skills,
            "breakdown":       brkdown,
            "job_match":       jd_sim,
            "matched_skills":  [s.title() for s in matched],
            "missing_skills":  [s.title() for s in missing],
            "suggestions":     suggestions,
            "resume_preview":  resume_text[:1000].strip(),
        })
    finally:
        try: os.remove(filepath)
        except: pass

@app.route("/download", methods=["POST"])
def download():
    data     = request.json or {}
    fmt      = data.pop("format","pdf").lower()
    template = data.pop("template","modern").lower()

    try:
        if fmt == "pdf":
            pdf_bytes = generate_pdf(data, template)
            return send_file(
                io.BytesIO(pdf_bytes),
                mimetype="application/pdf",
                as_attachment=True,
                download_name=f"resume_{template}.pdf"
            )
        else:
            docx_bytes = generate_docx(data)
            return send_file(
                io.BytesIO(docx_bytes),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name="resume.docx"
            )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True, port=5000)
